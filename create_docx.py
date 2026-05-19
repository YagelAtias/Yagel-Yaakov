import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

doc = Document()

def make_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    
    for run in p.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)

def add_paragraph(doc, text, style=None, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=False):
    # Split text by newlines so bullet points behave correctly in RTL
    lines = text.split('\n')
    first_p = None
    for line in lines:
        p = doc.add_paragraph(style=style)
        p.alignment = align
        run = p.add_run(line)
        if bold:
            run.bold = True
        make_rtl(p)
        if first_p is None:
            first_p = p
    return first_p

def add_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_heading(level=level)
    p.alignment = align
    p.add_run(text)
    make_rtl(p)
    return p

# Header Section
header_text = "המכללה האקדמית ספיר\nהמחלקה למדעי המחשב\nפרויקט גמר תכנון מפורט"
add_paragraph(doc, header_text, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

# Submitter details
submitters = "מגישים:\nיגל אטיאס\n[שם מגיש ב']\n\nמנחה: ד\"ר / מר / גב' [שם המנחה]"
add_paragraph(doc, submitters, align=WD_ALIGN_PARAGRAPH.LEFT)

# Title
add_heading(doc, "תכנון מפורט לפרויקט: Yagel-Yaakov (DistressEngine)", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)

# 1. הקדמה
add_heading(doc, "1. הקדמה: פסקה קצרה המתארת את הפרוייקט", level=2)
add_paragraph(doc, "מערכת Yagel-Yaakov (DistressEngine) היא פלטפורמה מוסדית לניהול מוסדות חינוך בתנאי פנימייה (ישיבות ואולפנות), המשלבת באופן אינטגרלי ניהול פדגוגי ולוגיסטי עם מנוע מתקדם לניתוח מצוקה רגשית. המערכת מספקת תמונה הוליסטית של התלמיד על ידי שילוב נתוני מערכת רגילים (ציונים, נוכחות, אישורי יציאה) יחד עם 'אותות' (Signals) סמויים כגון: קצב הקלדה (Typing Latency), עוצמת שמע (Acoustic Intensity), סמנטיקה קוגניטיבית ואנטרופיה. מטרתה לזהות דפוסי מצוקה בשלבים מוקדמים ולהתריע לצוות החינוכי בטרם יתפתחו למשבר.")

# 2. ארכיטקטורה
add_heading(doc, "2. ארכיטקטורה", level=2)
add_paragraph(doc, "המערכת בנויה בארכיטקטורת Client-Server (צד לקוח וצד שרת) קלאסית.\n\nצד לקוח (Client - React):\n- ממשק תלמיד (Student App)\n- ממשק מורה (Teacher Dashboard)\n- ממשק מנהל (Admin Dashboard)\n\nצד שרת (Server - FastAPI):\n- HTTP REST API (Auth & Routing)\n- מודול פדגוגי (Pedagogical)\n- מודול לוגיסטי/פנימייה (Leave Management)\n- מנוע ניתוח מצוקה (Distress Engine)\n- שכבת אבטחה והצפנה (Security & Encryption)\n\nבסיס נתונים:\n- Relational DB (SQLite / SQL)")

# 3. פרוט המודולים
add_heading(doc, "3. פרוט המודולים", level=2)
add_heading(doc, "מה רץ בצד הלקוח (Frontend):", level=3)
add_paragraph(doc, "צד הלקוח כתוב ב-React, ורץ על דפדפן המשתמש. הוא מכיל את ממשקי המשתמש (UI) השונים ואת הלוגיקה לאיסוף אותות (כגון זמני הקלדה והקלטת קול).\n"
"* סוג משתמש 1: תלמיד (Student) - רשאי להתחבר לממשק התלמיד. יכול להגיש בקשות יציאה, לענות על שאלונים או לכתוב ביומן. הממשק אוסף נתוני הקלדה ואודיו בזמן אמת ושולח אותם לשרת ללא הצגת הניתוח לתלמיד עצמו.\n"
"* סוג משתמש 2: מורה / איש צוות (Teacher/Counselor) - רשאי להתחבר לממשק המורה. הממשק כולל דשבורד המציג את מצב הכיתה, ציונים, בקשות יציאה הממתינות לאישור, וחשוב מכל - התראות מצוקה על תלמידים.\n"
"* סוג משתמש 3: מנהל מערכת (Admin) - רשאי לנהל את המוסד (הוספת מורים, הגדרת הרשאות, ניהול כיתות וכו'). למנהל יש ממשק ייעודי עם הרשאות נרחבות לכלל נתוני המוסד.")

add_heading(doc, "מה רץ בצד השרת (Backend):", level=3)
add_paragraph(doc, "צד השרת נכתב ב-Python באמצעות מסגרת FastAPI.\n"
"* מודול אימות והרשאות (Auth): מנהל כניסת משתמשים (Login) באמצעות טוקנים, ובודק איזה סוג משתמש מבקש לגשת למידע ומוודא הרשאות.\n"
"* מודול פדגוגי ולוגיסטי: מנהל את המידע המסורתי - ציונים, מערכת שעות, כיתות, ובקשות יציאה למחזור חיי פנימייה.\n"
"* מנוע ניתוח מצוקה (Distress Engine): הליבה של הפרויקט. מקבל נתונים מצד הלקוח ומריץ עליהם אלגוריתמים (מודלי שפה לחילוץ מרחק סמנטי, אלגוריתמי כיווץ לחישוב אנטרופיה, וניתוח קצב הקלדה). המודול מחשב ציון מצוקה משוקלל לכל תלמיד.\n"
"* מודול אבטחה והצפנה: פועל כ'שומר סף' למידע הרגיש. מוודא שנתוני טקסט גולמיים מוצפנים בצורה מאובטחת לפני שמירתם לבסיס הנתונים, תוך שימוש ב-Volatile Memory לעיבוד נתונים בזמן הריצה כדי למנוע דליפת מידע.")

# 4. בסיס נתונים
add_heading(doc, "4. בסיס נתונים", level=2)
add_paragraph(doc, "המערכת משתמשת בבסיס נתונים רלציוני (כגון SQLite). להלן סכמת הטבלאות המרכזיות וקשרי הגומלין (ERD):\n\n"
"1. טבלת Organizations (מוסדות)\n- id (PK, Integer), name (String)\n\n"
"2. טבלת Users (אנשי צוות)\n- id (PK, Integer), email, hashed_password, full_name, role, organization_id (FK), encrypted_dek, permissions\n\n"
"3. טבלת Classrooms (כיתות)\n- id (PK, Integer), name, organization_id (FK), teacher_id (FK)\n\n"
"4. טבלת Students (תלמידים)\n- id (PK, Integer), first_name, last_name, email, hashed_password, grade_level, organization_id (FK), classroom_id (FK)\n\n"
"5. טבלת DistressLog (יומן מצוקה / ניתוח אותות)\n- id (PK), student_id (FK), timestamp, encrypted_raw_text, overall_score, has_critical_alert, signals_metadata\n\n"
"6. טבלת DormLeave (בקשות יציאה / פנימייה)\n- id (PK), student_id (FK), leave_type, reason, destination, departure_date, return_date, is_approved, status, actual_return_time\n\n"
"7. טבלאות נוספות לניהול פדגוגי:\n- Grade (ציונים), BagrutStatus (בגרויות), Exam (מבחנים), ScheduleSlot (מערכת שעות)")

# 5. פרוטוקולים
add_heading(doc, "5. פרוטוקולים", level=2)
add_paragraph(doc, "המערכת משתמשת בפרוטוקולים נפוצים וסטנדרטיים, שנתמכים באופן טבעי (Native) על ידי הדפדפן וספריית צד השרת:\n\n"
"1. HTTP/HTTPS (REST API):\n"
"כל התקשורת בין ה-Frontend ל-Backend מבוססת על בקשות HTTP מסוג REST. המידע מועבר כולו בפורמט JSON. פרוטוקול זה נבחר עקב פשטות המימוש שלו ב-React (באמצעות fetch או axios) וב-FastAPI, שמייצרת ממשק REST באופן אוטומטי ואלגנטי.\n\n"
"2. Web API - MediaDevices:\n"
"לצורך איסוף האודיו, צד הלקוח משתמש ב-API הסטנדרטי של HTML5. המידע מוקלט או מעובד בצד הלקוח, ולאחר מכן הקובץ (או הטקסט המתומלל במידה ויש תמלול מקומי) נשלח דרך בקשת HTTP רגילה לשרת. אין צורך בתקשורת סוקטים בזמן אמת עקב שליחת המידע במנות (Batches).")

# 6. שונות/ אלגוריתמים
add_heading(doc, "6. שונות/ אלגוריתמים", level=2)
add_paragraph(doc, "מנוע הניתוח משתמש במספר אותות במקביל:\n"
"1. סמנטיקה קוגניטיבית: מודל שפה מבוסס Word2Vec (מודל הנתמך בשפה העברית) המחשב את המרחק הסמנטי (Word Mover's Distance) בין הטקסט של התלמיד למילון מונחי מצוקה קליניים.\n"
"2. אנטרופיה: שימוש באלגוריתמים ממשפחת Lempel-Ziv למדידת דחיסות הטקסט וזיהוי חזרתיות אובססיבית (Rumination).\n"
"3. Typing Latency: אלגוריתם המחשב את השונות וזמן ההשהייה בין הקשות כאינדיקציה ללחץ או חוסר יציבות רגשית.\n"
"4. Acoustic Intensity: מדידת חריגות עוצמת קול (דציבלים) המסמנים התפרצויות, בכי, או לחש סודי.\n\n"
"המידע מכלל האותות משוקלל לאינדקס מרכזי. אם האינדקס חוצה סף מסוים, המערכת מתעדת has_critical_alert=True ומציגה התראה בדשבורד המורה.")

output_path = r"c:\Users\yagel\Desktop\Coding\Python programming\Yagel-Yaakov\docs\System_Design_Sapir_v3.docx"
doc.save(output_path)
print(f"Document saved to {output_path} with STRICT RTL settings")
